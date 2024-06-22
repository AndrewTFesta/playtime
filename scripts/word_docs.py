"""
@title

@description

"""
import argparse
from pathlib import Path

from docx import Document

from playtime import project_properties

WORD_DIR = Path(project_properties.data_dir, 'word')


def document_text(doc_name):
    """
    This is a good start, it does not reflect text in tables, in headers, in footers and in foot-notes.

    https://stackoverflow.com/questions/25228106/how-to-extract-text-from-an-existing-docx-file-using-python-docx

    > python-docx will only find paragraphs and tables at the top-level of the document. In particular, paragraphs or tables "wrapped" in a "container" element will not be detected.
    > Most commonly, the "container" is a pending (not yet accepted) revision and this produces a similar behavior.
    > To extract the "wrapped" text, you'll need to know what the "wrapper" elements are. One way to do that is by dumping the XML of the document body:

    ```
    document = Document("my-document.docx")
    print(document._body._body.xml)
    ```

    > A paragraph element has a w:p tag, and you can inspect the output to look for those, some of which I expect will be inside another element.
    > Then you can extract those elements with XPath expressions, something like this, which would work if the "wrapper" element was <w:x>:

    ```
    from docx.text.paragraph import Paragraph

    body = document._body._body
    ps_under_xs = body.xpath("w:x//w:p")
    for p in ps_under_xs:
        paragraph = Paragraph(p, None)
        print(paragraph.text)
    ```

    > You could also just get all the <w:p> elements in the document, without regard to their "parentage" with something like this:

    ```
    ps = body.xpath(".//w:p")
    ```

    > The drawback of this is that some containers (like unaccepted revision marks) can contain text that has been "deleted" from the document, so you might get more than what you wanted.
    > In any case, this general approach should work for the job you've described. You can find more about XPath expressions on search if you need something more sophisticated.

    https://stackoverflow.com/questions/65412315/how-to-read-docx-originated-from-word-templates-with-python-docx

    :param doc_name:
    :return:
    """
    doc = Document(doc_name)
    text_lines = []
    for para in doc.paragraphs:
        text_lines.append(para.text)

    full_text = '\n'.join(text_lines)
    return full_text


def create_blank(doc_name):
    if not doc_name.parent.exists():
        doc_name.parent.mkdir(exist_ok=True, parents=True)

    document = Document()
    paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')
    prior_paragraph = paragraph.insert_paragraph_before('Lorem ipsum')
    document.add_heading('The REAL meaning of the universe')
    document.add_heading('The role of dolphins', level=2)
    document.add_page_break()
    document.save(doc_name)
    return


def main(main_args):
    """
    https://python-docx.readthedocs.io/en/latest/user/quickstart.html

    Docs are focused on creating and manipulating *existing* docs, don't explain much in the way of *reading* docs.
    :return:
    """
    create_name = Path(WORD_DIR, 'test_doc.docx')
    if not create_name.exists():
        create_blank(create_name)

    open_name = str(Path(WORD_DIR.parent, 'The Fantasy Realm', 'Immortal Guardian.docx'))
    text = document_text(open_name)
    print(text)
    return


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='')

    args = parser.parse_args()
    main(vars(args))
